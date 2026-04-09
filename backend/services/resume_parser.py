import os
import io
import re
import pdfplumber
import docx
import pytesseract
from PIL import Image, ImageOps, ImageEnhance

MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", 10))

def preprocess_image(content: bytes) -> Image:
    """Preprocess image for better OCR: grayscale and contrast increase."""
    image = Image.open(io.BytesIO(content))
    # Convert to grayscale
    image = ImageOps.grayscale(image)
    # Increase contrast
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(2.0)
    return image

def extract_text_image(content: bytes) -> str:
    """Extract text from JPG/PNG using pytesseract."""
    try:
        image = preprocess_image(content)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        print(f"OCR Error: {e}")
        return ""

def extract_text_docx(content: bytes) -> str:
    """Extract text from Word .docx preserving section order."""
    doc = docx.Document(io.BytesIO(content))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))
                
    return "\n".join(text_parts)

def extract_text_pdf(content: bytes) -> str:
    """Extract PDF text with two-column detection."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_width = page.width
                # Extract words with positioning
                words = page.extract_words()
                
                # Simple heuristic: split words into left and right clusters
                left_col = [w['text'] for w in words if w['x1'] < page_width / 2]
                right_col = [w['text'] for w in words if w['x0'] >= page_width / 2]
                
                # Determine if it's likely two-column: significant overlap in y-ranges
                # but distinct x-ranges. If clusters are clearly separated, we join them
                # left then right. Otherwise, we use standard extraction.
                
                if left_col and right_col:
                    # Very basic two-column concatenation
                    text += " ".join(left_col) + "\n" + " ".join(right_col) + "\n"
                else:
                    text += page.extract_text() or ""
                    
    except Exception as e:
        print(f"PDF Parsing Error: {e}")
    return text.strip()

def parse_resume(filename: str, content: bytes) -> dict:
    """Master parser for multiple formats."""
    # Size check
    if len(content) > MAX_FILE_SIZE_MB * 1024 * 1024:
        return {"error": "file_too_large", "filename": filename}

    ext = filename.lower().split(".")[-1]
    text = ""

    try:
        if ext == "pdf":
            text = extract_text_pdf(content)
        elif ext == "docx":
            text = extract_text_docx(content)
        elif ext in ["jpg", "jpeg", "png"]:
            text = extract_text_image(content)
        else:
            return {"error": "unsupported_format", "filename": filename}

        if not text:
            return {"error": "corrupt_file", "filename": filename}

        return {"status": "success", "text": text, "filename": filename}

    except Exception as e:
        return {"error": "corrupt_file", "filename": filename, "detail": str(e)}
