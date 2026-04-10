import io
import fitz  # PyMuPDF
import pdfplumber
import docx
import re
from google import genai
from google.genai import types
import os
from dotenv import load_dotenv

load_dotenv()


def _mime_for_image(ext: str) -> str:
    if ext in ("jpg", "jpeg"):
        return "image/jpeg"
    if ext == "png":
        return "image/png"
    return "image/png"


def extract_text_pdf_layout(content: bytes) -> str:
    """
    Layout-aware PDF extraction: group words into lines, split two-column rows with a separator
    so sidebars and body text are less likely to concatenate incorrectly.
    """
    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                w = page.width or 600
                mid = w * 0.52
                words = page.extract_words(keep_blank_chars=False, use_text_flow=False)
                if not words:
                    parts.append(page.extract_text() or "")
                    continue
                buckets: dict[float, list] = {}
                for word in words:
                    ymid = round(
                        (float(word.get("top", 0)) + float(word.get("bottom", 0))) / 2,
                        1,
                    )
                    buckets.setdefault(ymid, []).append(word)
                line_texts: list[str] = []
                for y in sorted(buckets.keys()):
                    row = sorted(buckets[y], key=lambda ww: float(ww["x0"]))
                    left = [ww for ww in row if float(ww["x0"]) < mid]
                    right = [ww for ww in row if float(ww["x0"]) >= mid]
                    if left and right:
                        line_texts.append(
                            " ".join(w["text"] for w in left)
                            + " | "
                            + " ".join(w["text"] for w in right)
                        )
                    else:
                        line_texts.append(" ".join(w["text"] for w in row))
                parts.append("\n".join(line_texts))
    except Exception:
        return ""
    return clean_text("\n\n".join(parts))


def extract_text_pdf(content: bytes) -> str:
    """Extract PDF text: layout-aware path first, then plumber default, then PyMuPDF."""
    layout_text = extract_text_pdf_layout(content)
    if len(layout_text.strip()) >= 80:
        return layout_text
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    except Exception:
        text = ""
    if len(text.strip()) < 40:
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            alt = "".join(page.get_text() for page in doc)
            doc.close()
            if len(alt.strip()) > len(text.strip()):
                text = alt
        except Exception:
            pass
    return clean_text(text)


def extract_text_docx(content: bytes) -> str:
    """Extract paragraphs and table cells from Word .docx."""
    doc = docx.Document(io.BytesIO(content))
    blocks: list[str] = []
    blocks.append("\n".join(p.text for p in doc.paragraphs if p.text.strip()))
    for table in doc.tables:
        rows_out = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_out.append(" | ".join(cells))
        if rows_out:
            blocks.append("\n".join(rows_out))
    return clean_text("\n\n".join(b for b in blocks if b))


async def extract_text_image_vision(content: bytes, ext: str) -> str:
    """Use Gemini Vision to OCR resume images."""
    api_key = os.getenv("GEMINI_API_KEY", "")
    if not api_key:
        return ""

    client = genai.Client(api_key=api_key)
    mime = _mime_for_image(ext)
    response = await client.aio.models.generate_content(
        model="gemini-1.5-flash",
        contents=[
            types.Part.from_bytes(data=content, mime_type=mime),
            "You are a senior HR specialist. Read this resume image and extract the full text accurately. "
            "Preserve section order (experience before education when visually above). "
            "Do not merge unrelated sidebar text into job bullets.",
        ],
    )
    return clean_text(response.text)


def clean_text(text: str) -> str:
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


async def parse_resume(filename: str, content: bytes) -> str:
    """Master pipeline for resume extraction based on file extension."""
    ext = filename.lower().split(".")[-1]

    if ext == "pdf":
        return extract_text_pdf(content)
    elif ext in ["docx", "doc"]:
        if ext == "doc":
            return ""
        return extract_text_docx(content)
    elif ext in ["jpg", "jpeg", "png"]:
        return await extract_text_image_vision(content, ext)
    else:
        return ""
